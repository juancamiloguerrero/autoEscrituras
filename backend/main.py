from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime
from typing import List
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os
import uuid

app = FastAPI()

# CORS para desarrollo
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Modelos de datos
class Persona(BaseModel):
    nombreCompleto: str
    identificacion: str
    expedicion: str
    domicilioDireccion: str
    departamentoResidencia: str
    ciudadResidencia: str
    telefono: str
    estadoCivil: str
    sexo: str
    correo: str
    ocupacion: str

class FormularioData(BaseModel):
    fechaOtorgamiento: str
    fechaPazYSalvo: str
    matriculaInmobiliaria: str
    cedulaCatastral: Optional[str]
    mayorExtension: bool
    departamento: str
    ciudad: str
    tipoPredio: str
    nombreLote: str
    direccion: str
    acto: str
    valorActo: float
    descripcionInmueble: Optional[str]
    oficinaInstrumentos: Optional[str]
    tradicion: Optional[str]
    complementosTradicion: Optional[str]
    viviendaFamiliar: str
    descripcionViviendaFamiliar: Optional[str]
    compradores: List[Persona]
    vendedores: List[Persona]

# Agrega esto justo después de las definiciones de modelos y antes de las funciones utilitarias
BASE_DIR = Path(__file__).parent.resolve()
TEMPLATE_PATH = BASE_DIR / "templates" / "formatoPlantilla.docx"

def formatear_parrafos(doc: Document):
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
        p.style.font.name = 'Arial'

def ajustar_margenes_superiores(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(3.1)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(2.3)
    section.gutter = Cm(0)


def fecha_a_texto(fecha_str: str) -> str:
    meses = [
        "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    dias_ordinales = {
        1: "primero", 2: "segundo", 3: "tercero", 4: "cuarto", 5: "quinto",
        6: "sexto", 7: "séptimo", 8: "octavo", 9: "noveno", 10: "décimo",
        11: "once", 12: "doce", 13: "trece", 14: "catorce", 15: "quince",
        16: "dieciséis", 17: "diecisiete", 18: "dieciocho", 19: "diecinueve", 20: "veinte",
        21: "veintiuno", 22: "veintidós", 23: "veintitrés", 24: "veinticuatro", 25: "veinticinco",
        26: "veintiséis", 27: "veintisiete", 28: "veintiocho", 29: "veintinueve", 30: "treinta", 31: "treinta y uno"
    }

    fecha = datetime.strptime(fecha_str, "%Y-%m-%d")
    dia = fecha.day
    mes = fecha.month
    anio = fecha.year

    dia_ordinal = dias_ordinales.get(dia, str(dia))
    mes_texto = meses[mes]
    anio_texto = f"{anio}"
    anio_extendido = "dos mil " + (f"veinticinco" if anio == 2025 else str(anio)[-2:])

    return f"a {dia_ordinal} ({dia:02}) día del mes de {mes_texto.capitalize()} del año {anio_extendido} ( {anio_texto} )"

def generar_parrafo_compradores(compradores: List[Persona]) -> str:
    parrafos = []
    for p in compradores:
        genero_domiciliado = "domiciliada" if p.sexo.lower() == "femenino" else "domiciliado"
        genero_identificada = "identificada" if p.sexo.lower() == "femenino" else "identificado"
        parrafo = (
            f"{p.nombreCompleto} mayor de edad, {genero_domiciliado} y residente en el Municipio de {p.ciudadResidencia} - {p.departamentoResidencia}, "
            f"en la {p.domicilioDireccion}, {genero_identificada} con la cédula de ciudadanía número {p.identificacion} expedida en {p.expedicion}, "
            f"de estado civil {p.estadoCivil},  teléfono {p.telefono}, ocupación {p.ocupacion}, correo electrónico {p.correo}"
        )
        parrafos.append(parrafo)

    introduccion = "PRIMER: Que transfiere a título de compraventa y enajenación real a favor de "
    cuerpo = "; ".join(parrafos)
    cierre = (
        ", conforme lo dispone el artículo 1506 del Código Civil (C.C). "
        "SIENDO LA VENTA, EL DERECHO DE DOMINIO, PROPIEDAD Y POSESIÓN, y todos los demás derechos reales que, "
        "junto con todas sus anexidades, usos, costumbres y servidumbres, tiene la vendedora sobre el total del cien por ciento (100%) "
        "del siguiente bien inmueble:"
    )

    return introduccion + cuerpo + cierre

def generar_parrafo_lote(datos: FormularioData) -> str:
    # Determinar cómo iniciar el párrafo según el tipo de vivienda
    if datos.viviendaFamiliar.lower() == "lote":
        inicio = "UN LOTE"
    elif datos.viviendaFamiliar.lower() == "otro":
        inicio = datos.descripcionViviendaFamiliar or ""
    else:
        inicio = datos.viviendaFamiliar  # fallback en caso de otros valores

    tipo_predio = datos.tipoPredio.upper()
    nombre_lote = datos.nombreLote.upper()
    ciudad = datos.ciudad
    departamento = datos.departamento

    return (
        f"{inicio} DE TERRENO {tipo_predio} DISTINGUIDO COMO {nombre_lote}, "
        f"Ubicado en la Vereda -----, del Municipio de {ciudad}, Departamento de {departamento},"
    )

def valor_a_texto(valor: float) -> str:
    entero = int(round(valor))
    return numero_a_letras(entero)

def generar_parrafo_segundo(datos: FormularioData) -> str:
    def obtener_articulo_y_rol(personas: List[Persona], rol_base: str):
        plural = len(personas) > 1
        sexos = {p.sexo.lower() for p in personas}

        if sexos == {"femenino"}:
            return ("las", rol_base + "as") if plural else ("la", rol_base + "a")
        elif sexos == {"masculino"}:
            return ("los", rol_base + "es") if plural else ("el", rol_base)
        else:
            return ("los", rol_base + "es")

    texto_valor = valor_a_texto(datos.valorActo)
    valor_formateado = f"${datos.valorActo:,.0f}".replace(",", ".")

    articulo_v, rol_v = obtener_articulo_y_rol(datos.vendedores, "vendedor")
    articulo_c, rol_c = obtener_articulo_y_rol(datos.compradores, "comprador")

    return (
        f"SEGUNDO: El precio de esta compraventa es por la cantidad de {texto_valor} "
        f"( {valor_formateado} ), moneda legal colombiana, dinero que {articulo_v} {rol_v} "
        f"declara haber recibido en efectivo de parte de {articulo_c} {rol_c} a entera satisfacción. "
    )

def generar_parrafo_compradoras_aceptan(datos: FormularioData) -> str:
    compradores = datos.compradores
    plural = len(compradores) > 1
    sexos = {p.sexo.lower() for p in compradores}

    # Determinar encabezado correcto según sexo y cantidad
    if sexos == {"femenino"}:
        encabezado = "las Compradoras" if plural else "la Compradora"
    elif sexos == {"masculino"}:
        encabezado = "los Compradores" if plural else "el Comprador"
    else:
        encabezado = "los Compradores" if plural else "el Comprador"

    # Obtener nombres completos
    nombres = ", ".join(p.nombreCompleto for p in compradores)

    # Construir párrafo
    return (
        f"Presente en este acto {encabezado}: {nombres}, de las condiciones civiles antes indicadas, quienes manifiestan: "
        "a) Que aceptan en todos y cada uno de sus términos la presente escritura y el contrato de compraventa que ella contiene a su favor; "
        "b) Que han cancelado la totalidad del valor acordado; y "
        "c) Que tienen recibido el inmueble a entera satisfacción, con todas sus mejoras y anexidades existentes. "
    )

def generar_parrafo_paragrafo_segundo(datos: FormularioData) -> str:
    vendedores = datos.vendedores
    plural = len(vendedores) > 1
    sexos = {p.sexo.lower() for p in vendedores}

    if sexos == {"femenino"}:
        encabezado = "LAS VENDEDORAS" if plural else "LA VENDEDORA"
    elif sexos == {"masculino"}:
        encabezado = "LOS VENDEDORES" if plural else "EL VENDEDOR"
    else:
        encabezado = "LOS VENDEDORES"

    return (
        f"PARAGRAFO SEGUNDO: {encabezado} O TRANSFERENTE DEJA EXPRESA CONSTANCIA BAJO LA GRAVEDAD DE JURAMENTO QUE "
        "SOBRE EL INMUEBLE QUE TRANSFIERE NO PESA PROTECCION ALGUNA QUE IMPIDA EL ACTO DE TRANSFERENCIA O ENAJENACION. "
    )

def generar_fichas_partes(datos: FormularioData) -> str:
    def encabezado_personas(lista: List[Persona], rol: str) -> str:
        plural = len(lista) > 1
        sexos = {p.sexo.lower() for p in lista}
        if sexos == {"femenino"}:
            return f"LAS {rol}AS" if plural else f"LA {rol}A"
        elif sexos == {"masculino"}:
            return f"LOS {rol}ES" if plural else f"EL {rol}"
        else:
            return f"LOS {rol}ES"

    def construir_ficha(persona: Persona, rol_extra: Optional[str] = "") -> str:
        rol_line = f"\n{rol_extra.upper()}" if rol_extra else ""
        return (
            f"{persona.nombreCompleto.upper()}                                                         HUELLA DACTILAR{rol_line}\n"
            f"DOCUMENTO DE IDENTIFICACION:                         INDICE DERECHO\n"
            f"TELEFONO O CELULAR:\n"
            f"DIRECCION:\n"
            f"CIUDAD:\n"
            f"E-MAIL:\n"
            f"ACTIVIDAD ECONOMICA:\n"
            f"ESTADO CIVIL:\n"
            f"PERSONA EXPUESTA POLITICAMENTE DECRETO 1674 DE 2016 SI__  NO __\n"
            f"CARGO:\n"
            f"FECHA DE VINCULACION:\n"
            f"FECHA DE DESVINCULACION:\n"
        )

    contenido = []

    # Vendedores
    encabezado_v = encabezado_personas(datos.vendedores, "VENDEDOR")
    contenido.append(f"{encabezado_v},\n")
    for v in datos.vendedores:
        contenido.append(construir_ficha(v))
    
    # Separador visual
    contenido.append("\n" + "\n")

    # Compradores
    encabezado_c = encabezado_personas(datos.compradores, "COMPRADOR")
    contenido.append(f"{encabezado_c},\n")
    for c in datos.compradores:
        ficha = construir_ficha(c)
        # Rol adicional si es estipulante
        if "estipulante" in c.ocupacion.lower():
            ficha = construir_ficha(c, "COMPRADORA Y ESTIPULANTE" if c.sexo.lower() == "femenino" else "COMPRADOR Y ESTIPULANTE")
        contenido.append(ficha)

    return "\n".join(contenido)


UNIDADES = (
    "", "UN", "DOS", "TRES", "CUATRO", "CINCO",
    "SEIS", "SIETE", "OCHO", "NUEVE", "DIEZ",
    "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE",
    "DIECISÉIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE", "VEINTE"
)

DECENAS = (
    "", "", "VEINTE", "TREINTA", "CUARENTA",
    "CINCUENTA", "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"
)

CENTENAS = (
    "", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
    "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"
)

def fecha_a_letras_estilo_mensual(fecha_str: str) -> str:
    meses = [
        "", "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
        "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"
    ]
    fecha = datetime.strptime(fecha_str, "%Y-%m-%d")
    mes = meses[fecha.month]
    dia = f"{fecha.day:02}"
    anio = fecha.year
    return f"{mes} {dia} DE {anio}"


def numero_a_letras(n: int) -> str:
    if n == 0:
        return "CERO PESOS MONEDA CORRIENTE"
    if n == 100:
        return "CIEN PESOS MONEDA CORRIENTE"

    resultado = ""
    millones = n // 1_000_000
    miles = (n % 1_000_000) // 1_000
    cientos = n % 1_000

    if millones > 0:
        if millones == 1:
            resultado += "UN MILLÓN "
        else:
            resultado += f"{convertir_trio(millones)} MILLONES "

    if miles > 0:
        if miles == 1:
            resultado += "MIL "
        else:
            resultado += f"{convertir_trio(miles)} MIL "

    if cientos > 0:
        resultado += convertir_trio(cientos)

    return resultado.strip() + " DE PESOS MONEDA CORRIENTE"

def convertir_trio(n: int) -> str:
    c = n // 100
    d = (n % 100) // 10
    u = n % 10

    texto = ""
    if n == 0:
        return ""

    if n <= 20:
        texto = UNIDADES[n]
    elif n < 100:
        texto = DECENAS[d]
        if u > 0:
            texto += f" Y {UNIDADES[u]}"
    else:
        texto = CENTENAS[c]
        resto = n % 100
        if resto > 0:
            texto += f" {convertir_trio(resto)}"

    return texto


def generar_parrafo_vendedores(vendedores: List[Persona]) -> str:
    parrafos = []
    for p in vendedores:
        genero_domiciliado = "domiciliada" if p.sexo.lower() == "femenino" else "domiciliado"
        genero_identificada = "identificada" if p.sexo.lower() == "femenino" else "identificado"
        parrafo = (
            f"{p.nombreCompleto} mayor de edad, {genero_domiciliado} y residente en el Municipio de {p.ciudadResidencia} - {p.departamentoResidencia}, "
            f"en la {p.domicilioDireccion}, {genero_identificada} con la cédula de ciudadanía número {p.identificacion} expedida en {p.expedicion}, "
            f"de estado civil {p.estadoCivil},  teléfono {p.telefono}, ocupación {p.ocupacion}, correo electrónico {p.correo}"
        )
        parrafos.append(parrafo)

    if len(parrafos) == 1:
        return parrafos[0] + ", manifestó:"
    else:
        return "; ".join(parrafos) + ", manifestaron:"


def reemplazar_campos(doc: Document, datos: FormularioData):
    valor_acto = f"${datos.valorActo:,.0f}".replace(",", ".")
    ciudad = datos.ciudad
    departamento = datos.departamento
    mayor_ext = "( MAYOR EXTENSION )" if datos.mayorExtension else ""
    direccion_completa = f"{datos.nombreLote}, {datos.direccion}"

    # Determinar encabezados dinámicos
    def encabezado_personas(lista, rol):
        plural = len(lista) > 1
        sexos = {p.sexo.lower() for p in lista}

        if sexos == {"femenino"}:
            return f"{rol}A" if not plural else f"{rol}AS", "femenino"
        elif sexos == {"masculino"}:
            return f"{rol}" if not plural else f"{rol}ES", "masculino"
        elif "femenino" in sexos and "masculino" in sexos:
            return f"{rol}ES", "mixto"
        else:
            return rol, "desconocido"

    def texto_personas(lista):
        bloques = []
        for p in lista:
            bloque = f"{p.nombreCompleto}\nC.C. {p.identificacion} expedida en {p.expedicion}"
            bloques.append(bloque)
        return "\n\n".join(bloques)

    reemplazos = {
        "fechaOtorgamiento": datos.fechaOtorgamiento,
        "fechaPazYSalvo": datos.fechaPazYSalvo,
        "matriculaInmobiliaria": datos.matriculaInmobiliaria,
        "cedulaCatastral": f"{datos.cedulaCatastral} {mayor_ext}",
        "mayorExtension": f"{mayor_ext} " if datos.mayorExtension else "",
        "ciudad": ciudad,
        "departamento": departamento,
        "tipoPredio": datos.tipoPredio,
        "nombreLote": datos.nombreLote,
        "direccion": datos.direccion,
        "valorActo": valor_acto,
        "codigoActo": datos.acto,
        "descripcionInmueble": datos.descripcionInmueble or "",
        "oficinaInstrumentos": datos.oficinaInstrumentos or "",
        "tradicion": datos.tradicion or "",
        "complementosTradicion": datos.complementosTradicion or "",
        "viviendaFamiliar": datos.viviendaFamiliar or "",
        "descripcionViviendaFamiliar": datos.descripcionViviendaFamiliar or "",
        "fechaOtorgamientoTexto": fecha_a_texto(datos.fechaOtorgamiento),
        "vendedoresParrafo": generar_parrafo_vendedores(datos.vendedores),
        "compradoresParrafo": generar_parrafo_compradores(datos.compradores),
        "loteParrafo": generar_parrafo_lote(datos),
        "parrafoSegundo": generar_parrafo_segundo(datos),
        "parrafoCompradoresAceptan": generar_parrafo_compradoras_aceptan(datos),
        "parrafoParagrafoSegundo": generar_parrafo_paragrafo_segundo(datos),
        "fichasPartes": generar_fichas_partes(datos),
        "letrasFecha": fecha_a_letras_estilo_mensual(datos.fechaOtorgamiento),
    }

    doc_text = "\n".join(p.text for p in doc.paragraphs)

    # Aplicar reemplazos simples
    for clave, valor in reemplazos.items():
        doc_text = doc_text.replace(f'{{“{clave}”}}', str(valor))

    # Insertar listas dinámicas
    encabezado_v, label_v = encabezado_personas(datos.vendedores, "VENDEDOR")
    encabezado_c, label_c = encabezado_personas(datos.compradores, "COMPRADOR")

    doc_text = doc_text.replace("\nIDENTIFICACIÓN VENDEDORA/VENDEDOR o VENDEDORAS/VENDEDORES", f"IDENTIFICACIÓN {encabezado_v}:\n\n")
    doc_text = doc_text.replace("\nIDENTIFICACION COMPRADORAS/COMPRADORES o COMPRADORA/COMPRADOR", f"IDENTIFICACIÓN {encabezado_c}:\n\n")
    doc_text = doc_text.replace("{vendedores}", texto_personas(datos.vendedores))
    doc_text = doc_text.replace("{compradores}", texto_personas(datos.compradores))
    doc_text = doc_text.replace("{vendedorEtiqueta}", encabezado_v)
    doc_text = doc_text.replace("{compradorEtiqueta}", encabezado_c)

    # Elimina todos los párrafos del documento original
    while len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    # Agregar cada línea del documento como párrafo independiente
    for bloque in doc_text.split("\n"):
        p = doc.add_paragraph(bloque.strip())
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)

    formatear_parrafos(doc)
    ajustar_margenes_superiores(doc)

# Endpoint principal
@app.post("/generar-doc/")
async def generar_documento(datos: FormularioData):
    try:
        # Verificación explícita de que existe la plantilla
        if not TEMPLATE_PATH.exists():
            error_detail = {
                "error": "Plantilla no encontrada",
                "ruta_esperada": str(TEMPLATE_PATH),
                "directorio_actual": str(BASE_DIR),
                "contenido_templates": os.listdir(str(BASE_DIR / "templates")) if os.path.exists(str(BASE_DIR / "templates")) else "Directorio no existe"
            }
            raise HTTPException(status_code=500, detail=error_detail)

        # Cargar documento
        doc = Document(str(TEMPLATE_PATH))
        
        # Procesar documento
        reemplazar_campos(doc, datos)

        # Guardar resultado
        output_id = str(uuid.uuid4())
        output_dir = BASE_DIR / "generated"
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / f"documento_{output_id}.docx"
        
        doc.save(str(output_path))

        return FileResponse(
            str(output_path),
            filename="escritura_generada.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except HTTPException:
        raise  # Re-lanzamos las excepciones HTTP que ya manejamos
    except Exception as e:
        error_detail = {
            "error": str(e),
            "type": type(e).__name__,
            "template_path": str(TEMPLATE_PATH),
            "exists": os.path.exists(str(TEMPLATE_PATH))
        }
        raise HTTPException(status_code=500, detail=error_detail)